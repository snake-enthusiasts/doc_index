import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class Main(object):
    """Main loop of application"""

    def open_doc(self, name):
        self.doc = docx.Document(name)

    def count_paragraphs(self):
        print(len(self.doc.paragraphs))

    def make_index_entry(self):
        self.MarkIndexEntry('Java', self.doc.paragraphs[3])

    def MarkIndexEntry(self, entry, paragraph):
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)

        run = paragraph.add_run()
        r = run._r
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' XE "%s" '%(entry)
        r.append(instrText)

        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

document = Main()
document.open_doc('testowy.docx')
document.count_paragraphs()
document.make_index_entry()
